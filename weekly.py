"""
weekly_report.py
-----------------
Single entry point for the weekly status report. Merges what used to be
3 separate files (chart_weekly.py, qwen_weekly_summary.py, docgen_weekly.py)
into one module that calls agent_weekely.py directly - so producing the
report is one function call:

    from weekly_report import generate_weekly_report
    filepath = generate_weekly_report(company_code="400", days_back=14)

Everything downstream (Streamlit button, scheduler job, email agent) just
calls generate_weekly_report() and gets back a ready .docx path.
"""

import datetime
import os

import matplotlib
matplotlib.use("Agg")  # no display needed, just save to file
import matplotlib.pyplot as plt
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import ollama

from agent_weekely import get_weekly_status_report


# =====================================================
# CONFIG
# =====================================================

OLLAMA_HOST = "http://68.178.160.26:11434"
QWEN_MODEL = "qwen2.5:7b"

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # dark blue, matches docgen.py
HEADER_FILL = "1F4E79"

BAR_COLORS = {
    "Completed": "#2E7D32",   # green
    "InProcess": "#1565C0",   # blue
    "Delay": "#C62828",       # red
}


# =====================================================
# 1. CHART
# =====================================================

def _build_status_totals_chart(totals, output_path="weekly_status_chart.png"):
    labels = ["Completed", "InProcess", "Delay"]
    values = [
        totals.get("completed_total", 0),
        totals.get("inprocess_total", 0),
        totals.get("delay_total", 0),
    ]
    colors = [BAR_COLORS[label] for label in labels]

    fig, ax = plt.subplots(figsize=(6, 2.6), dpi=150)
    bars = ax.bar(labels, values, color=colors, width=0.55)

    for bar, value in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            str(value),
            ha="center",
            va="bottom",
            fontsize=11,
            fontweight="bold",
        )

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.get_yaxis().set_visible(False)
    ax.tick_params(axis="x", labelsize=11)

    max_val = max(values) if values else 0
    ax.set_ylim(0, max_val * 1.25 if max_val else 1)

    fig.tight_layout()
    fig.savefig(output_path, transparent=True)
    plt.close(fig)

    return output_path


# =====================================================
# 2. QWEN SUMMARY
# =====================================================

def _summarize_weekly_totals(totals, delay_summary_df=None):
    """
    Only ever fed already-computed numbers - never raw rows - so it can't
    hallucinate counts, only phrase them. Falls back to a plain factual
    sentence if the Ollama call fails.
    """
    client = ollama.Client(host=OLLAMA_HOST)

    top_delay_line = ""
    if delay_summary_df is not None and not delay_summary_df.empty:
        worst = delay_summary_df.sort_values("TotalCount", ascending=False).iloc[0]
        top_delay_line = (
            f"The single highest delay count belongs to {worst['Responsible']} "
            f"under manager {worst['Manager']}, with {worst['TotalCount']} delayed items."
        )

    days_back = totals.get("days_back", 7)
    completed_line = f"{totals.get('completed_total', 0)} items completed in the last {days_back} days."
    if totals.get("completed_total", 0) == 0 and totals.get("last_completed_fallback"):
        fb = totals["last_completed_fallback"]
        completed_date = fb["CompletedOn"].strftime("%Y-%m-%d") if fb.get("CompletedOn") else "unknown date"
        completed_line = (
            f"No items completed in the last {days_back} days. "
            f"Most recent completion was '{fb['ItemName']}' on {completed_date}."
        )

    prompt = f"""You are writing a single short summary line (max 2 sentences) for a weekly
company project-status report. Use ONLY the facts below - do not invent numbers,
names, or dates that are not given. Be direct and factual, no fluff.

Facts:
- {completed_line}
- {totals.get('inprocess_total', 0)} items currently in process.
- {totals.get('delay_total', 0)} items currently delayed (overdue).
- {top_delay_line}

Write the summary now:"""

    try:
        response = client.chat(
            model=QWEN_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options={"num_thread": 32, "num_predict": 150, "num_ctx": 2048},
            keep_alive="30m",
        )
        return response["message"]["content"].strip()

    except Exception as e:
        print(f"Qwen summary generation failed: {e}")
        return (
            f"{completed_line} {totals.get('inprocess_total', 0)} items in process, "
            f"{totals.get('delay_total', 0)} delayed."
        )


# =====================================================
# 3. DOCX RENDERING
# =====================================================

def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_summary_table(doc, title, df, note=""):
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT_COLOR

    if note:
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(note)
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if df is None or df.empty:
        empty_p = doc.add_paragraph()
        empty_run = empty_p.add_run("No items in this category.")
        empty_run.italic = True
        empty_run.font.size = Pt(9)
        doc.add_paragraph()
        return

    columns = ["Manager", "Responsible", "Status", "TotalCount"]
    headers = ["Manager", "Employee", "Status", "Count"]

    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
        _shade_cell(cell, HEADER_FILL)

    manager_start_row = 1
    previous_manager = None
    for row_number, (_, row) in enumerate(df.iterrows(), start=1):
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            _set_cell_text(cells[i], row[col], size=9)

        manager = str(row["Manager"])
        if previous_manager is not None and manager != previous_manager:
            if row_number - 1 > manager_start_row:
                merged_cell = table.cell(manager_start_row, 0).merge(table.cell(row_number - 1, 0))
                _set_cell_text(merged_cell, previous_manager, size=9)
                merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            manager_start_row = row_number
        previous_manager = manager

    # Merge adjacent Manager cells so a manager appears once beside all of
    # their employees, while each employee keeps their own count row.
    if len(df) > 1 and len(df) > manager_start_row:
        merged_cell = table.cell(manager_start_row, 0).merge(table.cell(len(df), 0))
        _set_cell_text(merged_cell, previous_manager, size=9)
        merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()


def _completed_note(totals):
    days_back = totals.get("days_back", 7)
    completed_count = totals.get("completed_action_total", totals.get("completed_total", 0))
    if completed_count == 0 and totals.get("last_completed_fallback"):
        fb = totals["last_completed_fallback"]
        completed_date = (
            fb["CompletedOn"].strftime("%Y-%m-%d") if fb.get("CompletedOn") else "unknown date"
        )
        return f"None completed in this window. Last completed: '{fb['ItemName']}' on {completed_date}."
    return f"Completed within the last {days_back} days."


def _add_employee_count_breakdown(doc, title, df, item_label):
    """Add one manager heading with employee-level item counts beneath it."""
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR

    if df.empty:
        doc.add_paragraph(f"No {item_label.lower()} in this category.", style="List Bullet")
        return

    singular_label = item_label[:-1] if item_label.endswith("s") else item_label

    employee_counts = (
        df.groupby(["Manager", "Responsible"], dropna=False)
        .size()
        .reset_index(name="Count")
        .sort_values(["Manager", "Responsible"])
    )
    for manager, manager_rows in employee_counts.groupby("Manager", sort=False):
        doc.add_paragraph(str(manager) or "Unassigned manager", style="List Bullet")
        for _, row in manager_rows.iterrows():
            doc.add_paragraph(
                f"{row['Responsible'] or 'Unassigned employee'}: {row['Count']} "
                f"{singular_label if row['Count'] == 1 else item_label}",
                style="List Bullet 2",
            )


def _summarize_employee_item_counts(df):
    """Return one row per manager, employee, and status with a real row count."""
    columns = ["Section", "Manager", "TmId", "Responsible", "Status", "TotalCount"]
    if df.empty:
        return pd.DataFrame(columns=columns)

    return (
        df.groupby(["Section", "Manager", "TmId", "Responsible", "Status"], dropna=False)
        .size()
        .reset_index(name="TotalCount")
        .sort_values(["Manager", "Responsible"])
        .reset_index(drop=True)
    )


def _prepare_weekly_report_data(company_code=None, days_back=7):
    completed_df, delay_df, inprocess_df, totals = get_weekly_status_report(
        company_code=company_code, days_back=days_back
    )

    # Completion activity is recorded on Actions, while the open-status
    # sections report parent-project status.
    completed_actions_only = completed_df[completed_df["RowType"] == "Action"]
    delay_projects_only = delay_df[delay_df["RowType"] == "Project"]
    inprocess_projects_only = inprocess_df[inprocess_df["RowType"] == "Project"]
    totals["completed_action_total"] = len(completed_actions_only)
    totals["delay_project_total"] = len(delay_projects_only)
    totals["inprocess_project_total"] = len(inprocess_projects_only)

    return (
        totals,
        _summarize_employee_item_counts(completed_actions_only),
        _summarize_employee_item_counts(delay_projects_only),
        _summarize_employee_item_counts(inprocess_projects_only),
    )


def _weekly_status_points(totals):
    days_back = totals.get("days_back", 7)
    points = [
        f"Completed: {totals.get('completed_action_total', 0)} actions completed in the last {days_back} days.",
        f"Delay: {totals.get('delay_project_total', 0)} projects currently overdue.",
        f"In Process / Open: {totals.get('inprocess_project_total', 0)} projects currently active, on hold, cancelled, or awaiting approval.",
    ]
    if totals.get("completed_action_total", 0) == 0 and totals.get("last_completed_fallback"):
        fallback = totals["last_completed_fallback"]
        completed_date = fallback["CompletedOn"].strftime("%Y-%m-%d") if fallback.get("CompletedOn") else "unknown date"
        points[0] = (
            f"Completed: no actions completed in the last {days_back} days. "
            f"Last completed project: '{fallback['ItemName']}' on {completed_date}."
        )
    return points


def _build_weekly_report_docx(
    filepath,
    company_code,
    totals,
    completed_summary_df,
    delay_summary_df,
    inprocess_summary_df,
):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"Weekly Project Status \u2014 Company {company_code}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    window_days = totals.get("days_back", 7)
    sub_run = subtitle_p.add_run(
        f"Generated {datetime.date.today().isoformat()} \u2014 Completed window: last {window_days} days"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    overview = doc.add_paragraph()
    overview_run = overview.add_run("Status at a glance")
    overview_run.bold = True
    overview_run.font.size = Pt(13)
    overview_run.font.color.rgb = ACCENT_COLOR

    for point in _weekly_status_points(totals):
        doc.add_paragraph(point, style="List Bullet")

    doc.add_paragraph()
    _add_summary_table(doc, "\u2705 Completed Actions", completed_summary_df,
                       note=_completed_note(totals))
    _add_summary_table(doc, "\U0001F534 Delay (Overdue) Projects", delay_summary_df,
                       note="Projects past their deadline and still In Process.")
    _add_summary_table(doc, "\U0001F535 In Process / Open Projects", inprocess_summary_df,
                       note="Projects currently active, on hold, cancelled, or awaiting approval.")

    doc.save(filepath)
    return filepath


def _add_pdf_summary_table(elements, title, df, note, styles):
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    elements.append(Paragraph(title, styles["SectionHeading"]))
    if note:
        elements.append(Paragraph(note, styles["Note"]))

    if df is None or df.empty:
        elements.append(Paragraph("No items in this category.", styles["Empty"]))
        elements.append(Spacer(1, 0.16 * inch))
        return

    headers = ["Manager", "Employee", "Status", "Count"]
    table_data = [headers]
    previous_manager = None
    for _, row in df.iterrows():
        manager = str(row["Manager"])
        table_data.append([
            manager if manager != previous_manager else "",
            str(row["Responsible"]),
            str(row["Status"]),
            str(row["TotalCount"]),
        ])
        previous_manager = manager

    table = Table(table_data, colWidths=[1.7 * inch, 2.0 * inch, 1.7 * inch, 0.7 * inch], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB3C2")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (3, 1), (3, -1), "RIGHT"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FB")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 0.18 * inch))


def _build_weekly_report_pdf(
    filepath,
    company_code,
    totals,
    completed_summary_df,
    delay_summary_df,
    inprocess_summary_df,
):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ReportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=1,
        textColor=colors.HexColor("#1F4E79"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="Subtitle",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        alignment=1,
        textColor=colors.HexColor("#606060"),
        spaceAfter=14,
    ))
    styles.add(ParagraphStyle(
        name="SectionHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#1F4E79"),
        spaceBefore=8,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="Point",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
        leftIndent=12,
        bulletIndent=2,
        spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name="Note",
        parent=styles["Normal"],
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#606060"),
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="Empty",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#606060"),
    ))

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    elements = [
        Paragraph(f"Weekly Project Status - Company {company_code}", styles["ReportTitle"]),
        Paragraph(
            f"Generated {datetime.date.today().isoformat()} - Completed window: last {totals.get('days_back', 7)} days",
            styles["Subtitle"],
        ),
        Paragraph("Status at a glance", styles["SectionHeading"]),
    ]

    for point in _weekly_status_points(totals):
        elements.append(Paragraph(point, styles["Point"], bulletText="-"))

    elements.append(Spacer(1, 0.08 * inch))
    _add_pdf_summary_table(elements, "Completed Actions", completed_summary_df, _completed_note(totals), styles)
    _add_pdf_summary_table(elements, "Delay (Overdue) Projects", delay_summary_df, "Projects past their deadline and still In Process.", styles)
    _add_pdf_summary_table(elements, "In Process / Open Projects", inprocess_summary_df, "Projects currently active, on hold, cancelled, or awaiting approval.", styles)

    doc.build(elements)
    return filepath


# =====================================================
# 4. TOP-LEVEL ENTRY POINT
# =====================================================

def generate_weekly_report(company_code=None, days_back=7, output_dir="."):
    """
    Full pipeline in one call: pulls data via agent_weekely, renders the
    concise status-summary docx, and returns the file path.

        filepath = generate_weekly_report(company_code="400", days_back=14)
    """
    os.makedirs(output_dir, exist_ok=True)
    totals, completed_summary, delay_summary, inprocess_summary = _prepare_weekly_report_data(
        company_code=company_code, days_back=days_back
    )

    docx_path = f"{output_dir}/Weekly_Status_Report_{company_code or 'all'}.docx"
    _build_weekly_report_docx(
        filepath=docx_path,
        company_code=company_code or "All",
        totals=totals,
        completed_summary_df=completed_summary,
        delay_summary_df=delay_summary,
        inprocess_summary_df=inprocess_summary,
    )

    return docx_path


def generate_weekly_report_pdf(company_code=None, days_back=7, output_dir="output/pdf"):
    """
    Pulls the same weekly report data and renders a browser-friendly PDF for
    Streamlit preview/download.
    """
    os.makedirs(output_dir, exist_ok=True)
    totals, completed_summary, delay_summary, inprocess_summary = _prepare_weekly_report_data(
        company_code=company_code, days_back=days_back
    )

    pdf_path = f"{output_dir}/Weekly_Status_Report_{company_code or 'all'}.pdf"
    _build_weekly_report_pdf(
        filepath=pdf_path,
        company_code=company_code or "All",
        totals=totals,
        completed_summary_df=completed_summary,
        delay_summary_df=delay_summary,
        inprocess_summary_df=inprocess_summary,
    )

    return pdf_path


if __name__ == "__main__":
    # Quick manual test: python weekly_report.py
    path = generate_weekly_report(company_code="400", days_back=14)
    print(f"Report saved to {path}")
